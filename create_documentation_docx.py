
import re
from docx import Document
from docx.shared import Pt

# Initialize Document
doc = Document()
doc.add_heading('EduSync: Advanced Quiz Management System', 0)

# Read Markdown File
md_path = r"C:\Users\Shani\.gemini\antigravity\brain\15e47281-694d-4f3a-b3d0-f5945771af85\project_documentation.md"
try:
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
except FileNotFoundError:
    print(f"Error: Could not find file at {md_path}")
    exit(1)

def add_paragraph_with_formatting(paragraph, text):
    # Regex for bold: **text**
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)

for line in lines:
    line = line.strip()
    if not line:
        continue
    
    if line.startswith('# '):
        # Already used title, but maybe add as H1 if duplicates exist,
        # but the first line is usually the title. We'll skip the main title if it matches.
        if "EduSync: Advanced Quiz Management System" in line:
            continue
        doc.add_heading(line[2:], level=1)
    elif line.startswith('## '):
        doc.add_heading(line[3:], level=2)
    elif line.startswith('### '):
        doc.add_heading(line[4:], level=3)
    elif line.startswith('* '):
        p = doc.add_paragraph(style='List Bullet')
        add_paragraph_with_formatting(p, line[2:])
    elif line.startswith('    * '): # Indented list
        p = doc.add_paragraph(style='List Bullet 2') # Might need custom style or just List Bullet
        p.paragraph_format.left_indent = Pt(20)
        add_paragraph_with_formatting(p, line[6:])
    else:
        p = doc.add_paragraph()
        add_paragraph_with_formatting(p, line)

start_output_path = r"C:\Users\Shani\Documents\GitHub\quiz app flutter\EduSync_Project_Documentation.docx"
doc.save(start_output_path)
print(f"Document saved to: {start_output_path}")
